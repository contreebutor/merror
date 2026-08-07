"""Text extraction from uploaded documents.

Parsing happens entirely in-process — no document content is sent anywhere.
"""

from __future__ import annotations

import io
import logging

logger = logging.getLogger("merror.documents")

# Extension -> canonical kind. Content type from the browser is unreliable
# (Safari and Windows disagree about .docx), so the extension decides.
SUPPORTED_EXTENSIONS = {
    ".txt": "text",
    ".md": "text",
    ".markdown": "text",
    ".pdf": "pdf",
    ".docx": "docx",
}

MAX_UPLOAD_BYTES = 25 * 1024 * 1024  # 25 MB


class UnsupportedDocumentError(ValueError):
    """The file extension is not one we can parse."""


class DocumentParseError(ValueError):
    """The file matched a supported type but could not be read."""


def extension_of(filename: str) -> str:
    """Lowercased extension including the dot, or '' if there is none."""
    _, dot, ext = filename.rpartition(".")
    return f".{ext.lower()}" if dot else ""


def is_supported(filename: str) -> bool:
    return extension_of(filename) in SUPPORTED_EXTENSIONS


def _extract_txt(data: bytes) -> str:
    """Decode plain text, tolerating non-UTF-8 encodings."""
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    # latin-1 cannot actually fail, but stay defensive rather than raise.
    return data.decode("utf-8", errors="replace")


def _extract_pdf(data: bytes) -> str:
    """Pull text from a PDF, page by page.

    Scanned PDFs hold images rather than text and yield nothing here; the caller
    reports that as an empty extraction rather than silently storing a blank
    memory.
    """
    from pypdf import PdfReader
    from pypdf.errors import PdfReadError

    try:
        reader = PdfReader(io.BytesIO(data))
    except (PdfReadError, OSError, ValueError) as exc:
        raise DocumentParseError(f"Could not read PDF: {exc}") from exc

    if reader.is_encrypted:
        # An empty user password is common for "owner-locked" PDFs.
        try:
            if reader.decrypt("") == 0:
                raise DocumentParseError(
                    "This PDF is password protected. Remove the password and retry."
                )
        except (NotImplementedError, PdfReadError) as exc:
            raise DocumentParseError(f"Could not decrypt PDF: {exc}") from exc

    pages: list[str] = []
    for number, page in enumerate(reader.pages, start=1):
        try:
            pages.append(page.extract_text() or "")
        except Exception as exc:  # noqa: BLE001 - one bad page must not kill the file
            logger.warning("Skipping unreadable page %d: %s", number, exc)

    return "\n\n".join(p.strip() for p in pages if p.strip())


def _extract_docx(data: bytes) -> str:
    """Pull text from a .docx, including table cells.

    Legacy .doc is a different, binary format that python-docx cannot read; it
    is rejected by extension before reaching here.
    """
    import zipfile

    import docx
    from docx.opc.exceptions import PackageNotFoundError

    try:
        document = docx.Document(io.BytesIO(data))
    except (
        # A .docx is a zip archive; a corrupt or misnamed file surfaces as
        # BadZipFile, which is not a ValueError and would otherwise escape.
        zipfile.BadZipFile,
        PackageNotFoundError,
        KeyError,
        ValueError,
    ) as exc:
        raise DocumentParseError(
            f"Could not read .docx file: {exc}. If this is a legacy .doc, "
            "re-save it as .docx first."
        ) from exc

    blocks = [p.text.strip() for p in document.paragraphs if p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))

    return "\n\n".join(blocks)


def extract_text(filename: str, data: bytes) -> str:
    """Extract plain text from an uploaded document.

    Raises UnsupportedDocumentError for unknown types and DocumentParseError
    when a supported type cannot be read.
    """
    extension = extension_of(filename)
    kind = SUPPORTED_EXTENSIONS.get(extension)
    if kind is None:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise UnsupportedDocumentError(
            f"Cannot read '{extension or filename}'. Supported types: {supported}."
        )

    if not data:
        raise DocumentParseError("The uploaded file is empty.")

    extractor = {"text": _extract_txt, "pdf": _extract_pdf, "docx": _extract_docx}[kind]
    text = extractor(data)

    logger.info("Extracted %d chars from %s (%s)", len(text), filename, kind)
    return text
