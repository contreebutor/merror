"""Builders for real files and fake API responses used in tests.

Document and image bytes are genuine rather than mocked, so the parsers and
sniffers are exercised the same way they will be in production. Only the
Anthropic API call itself is faked — it costs money and needs a network.
"""

import io
import struct
import zlib
from dataclasses import dataclass, field


def make_pdf(lines: list[str]) -> bytes:
    """Build a minimal single-page PDF containing `lines` as text."""
    content = (
        "BT /F1 12 Tf 72 720 Td 14 TL\n"
        + "\n".join(f"({line}) Tj T*" for line in lines)
        + "\nET"
    )
    objects = [
        "<</Type/Catalog/Pages 2 0 R>>",
        "<</Type/Pages/Kids[3 0 R]/Count 1>>",
        "<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]"
        "/Contents 4 0 R/Resources<</Font<</F1 5 0 R>>>>>>",
        f"<</Length {len(content)}>>\nstream\n{content}\nendstream",
        "<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>",
    ]

    out = "%PDF-1.4\n"
    offsets = []
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{index} 0 obj\n{obj}\nendobj\n"

    xref_offset = len(out)
    out += f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n"
    out += "".join(f"{offset:010d} 00000 n \n" for offset in offsets)
    out += (
        f"trailer\n<</Size {len(objects) + 1}/Root 1 0 R>>\n"
        f"startxref\n{xref_offset}\n%%EOF"
    )
    return out.encode("latin-1")


def make_empty_pdf() -> bytes:
    """A structurally valid PDF with no text layer, like a scan."""
    return make_pdf([])


def make_docx(paragraphs: list[str], table_rows: list[list[str]] | None = None) -> bytes:
    """Build a real .docx file with the given paragraphs and optional table."""
    import docx

    document = docx.Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)

    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for row_index, row in enumerate(table_rows):
            for cell_index, value in enumerate(row):
                table.cell(row_index, cell_index).text = value

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# --- Real image bytes ------------------------------------------------------


def make_png(width: int = 1, height: int = 1) -> bytes:
    """Build a minimal valid PNG."""

    def chunk(tag: bytes, payload: bytes) -> bytes:
        return (
            struct.pack(">I", len(payload))
            + tag
            + payload
            + struct.pack(">I", zlib.crc32(tag + payload) & 0xFFFFFFFF)
        )

    header = struct.pack(">2I5B", width, height, 8, 2, 0, 0, 0)
    raw = b"".join(b"\x00" + b"\xff\x00\x00" * width for _ in range(height))
    return (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", header)
        + chunk(b"IDAT", zlib.compress(raw))
        + chunk(b"IEND", b"")
    )


def make_jpeg() -> bytes:
    """Build a byte string with a valid JPEG signature."""
    return b"\xff\xd8\xff\xe0" + b"\x00\x10JFIF\x00" + b"\x00" * 32 + b"\xff\xd9"


def make_gif() -> bytes:
    """Build a byte string with a valid GIF signature."""
    return b"GIF89a" + b"\x01\x00\x01\x00\x00\x00\x00" + b"\x3b"


def make_webp() -> bytes:
    """Build a byte string with a valid RIFF/WEBP signature."""
    body = b"VP8 " + b"\x00" * 16
    return b"RIFF" + struct.pack("<I", len(body) + 4) + b"WEBP" + body


# --- Fake Anthropic responses ----------------------------------------------


@dataclass
class FakeTextBlock:
    text: str
    type: str = "text"


@dataclass
class FakeStopDetails:
    category: str | None = None
    explanation: str = ""


@dataclass
class FakeMessage:
    """Stand-in for anthropic.types.Message, with only the fields we read."""

    content: list = field(default_factory=list)
    stop_reason: str = "end_turn"
    stop_details: FakeStopDetails | None = None


class FakeMessages:
    def __init__(self, response):
        self._response = response
        self.calls: list[dict] = []

    def create(self, **kwargs):
        self.calls.append(kwargs)
        if isinstance(self._response, Exception):
            raise self._response
        return self._response


class FakeAnthropic:
    """Minimal stand-in for anthropic.Anthropic."""

    def __init__(self, response):
        self.messages = FakeMessages(response)


def text_response(text: str) -> FakeMessage:
    return FakeMessage(content=[FakeTextBlock(text=text)])


def refusal_response(category: str = "privacy") -> FakeMessage:
    return FakeMessage(
        content=[],
        stop_reason="refusal",
        stop_details=FakeStopDetails(category=category),
    )
